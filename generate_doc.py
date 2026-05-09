from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Helpers ──────────────────────────────────────────────────────────────────
def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_col_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(width_cm * 567)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)

def h1(doc, text):
    p = doc.add_heading(text, level=1)
    run = p.runs[0]
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    run.font.size = Pt(18)

def h2(doc, text):
    p = doc.add_heading(text, level=2)
    run = p.runs[0]
    run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
    run.font.size = Pt(14)

def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(level * 0.5)

def add_table(doc, headers, rows, col_widths=None, header_color="2563EB"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        shade_cell(cell, header_color)
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if col_widths:
            set_col_width(cell, col_widths[i])
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            if col_widths:
                set_col_width(cell, col_widths[c_idx])
            if r_idx % 2 == 0:
                shade_cell(cell, "F8FAFC")
    doc.add_paragraph()
    return table

# ════════════════════════════════════════════════════════════════════════════
# PAGE DE TITRE
# ════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("UniSchedule")
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Plateforme de Gestion des Emplois du Temps Universitaires")
r2.font.size = Pt(16)
r2.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_str = datetime.date.today().strftime("%d/%m/%Y")
r3 = p3.add_run("Documentation Technique & Fonctionnelle  -  " + date_str)
r3.font.size = Pt(11)
r3.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 1. VUE D'ENSEMBLE
# ════════════════════════════════════════════════════════════════════════════
h1(doc, "1. Vue d'ensemble du projet")
doc.add_paragraph(
    "UniSchedule est une application web de gestion des emplois du temps universitaires. "
    "Elle couvre l'ensemble du cycle academique : creation de la structure "
    "(universite -> faculte -> departement -> promotion -> classe), "
    "planification des cours, gestion des professeurs, suivi des etudiants, "
    "depot et correction des TPs, publication des notes et bulletins."
)

h2(doc, "1.1 Stack technique")
add_table(doc,
    ["Composant", "Technologie", "Role"],
    [
        ["Frontend / UI",    "Streamlit 1.x",                          "Interface web multi-pages, rendu HTML/CSS"],
        ["Base de donnees",  "PostgreSQL 14+",                         "Stockage relationnel complet"],
        ["Connecteur DB",    "psycopg2 + ThreadedConnectionPool",       "Pool de connexions (configurable)"],
        ["Stockage fichiers","Supabase Storage",                        "PDFs, logos, pieces jointes"],
        ["Authentification", "bcrypt + session_state",                  "Hachage mots de passe, sessions"],
        ["Export PDF",       "ReportLab / WeasyPrint",                  "Bulletins et releves de notes"],
        ["Export iCal",      "icalendar",                               "Export emploi du temps format .ics"],
        ["Serveur",          "VPS Linux recommande",                    "Nginx + systemd + SSL Let's Encrypt"],
    ],
    col_widths=[4, 5, 6.5]
)

h2(doc, "1.2 Structure des fichiers")
add_table(doc,
    ["Fichier / Dossier", "Description"],
    [
        ["app.py",                         "Point d'entree - navigation multi-roles, sidebar"],
        ["db/connection.py",               "Pool de connexions PostgreSQL, execute_query()"],
        ["db/queries.py",                  "Toutes les requetes SQL (30+ classes de queries)"],
        ["utils/auth.py",                  "Authentification, hachage, creation comptes"],
        ["utils/components.py",            "Composants UI reutilisables (CSS global, cartes, tableaux)"],
        ["utils/storage.py",               "Upload/download fichiers Supabase Storage"],
        ["utils/pdf_export.py",            "Generation bulletins PDF"],
        ["utils/ical_export.py",           "Export emploi du temps .ics"],
        ["utils/notifications.py",         "Notifications in-app"],
        ["pages/1_Accueil.py",             "Page publique - liste des universites, stats plateforme"],
        ["pages/2_Horaire.py",             "Page publique - emploi du temps par classe"],
        ["pages/7_Admin_Login.py",         "Connexion administrateurs (toute hierarchie)"],
        ["pages/8_Admin_Dashboard.py",     "Dashboard admin (super / universite / faculte / departement)"],
        ["pages/9_Prof_Dashboard.py",      "Dashboard professeur (8 onglets)"],
        ["pages/10_Student_Auth.py",       "Connexion et inscription etudiant"],
        ["pages/11_Student_Dashboard.py",  "Dashboard etudiant (10 onglets)"],
        ["pages/12_Prof_Auth.py",          "Connexion professeur"],
        [".streamlit/secrets.toml",        "Configuration DB et Supabase (ne pas versionner)"],
        ["migration_v*.sql",               "Scripts de migration base de donnees (v1 a v30)"],
    ],
    col_widths=[5.5, 10]
)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 2. HIERARCHIE ET ROLES
# ════════════════════════════════════════════════════════════════════════════
h1(doc, "2. Hierarchie des roles")
doc.add_paragraph(
    "L'application implemente une hierarchie a 6 niveaux. "
    "Chaque role a un perimetre d'action strictement delimite."
)
add_table(doc,
    ["Role", "Perimetre", "Cree par"],
    [
        ["super_admin",        "Toutes les universites de la plateforme",              "Systeme"],
        ["admin_universite",   "Une universite (profs, facultes, annonces)",            "Super Admin"],
        ["admin_faculte",      "Une faculte (departements, affiliations profs)",         "Admin Universite"],
        ["admin_departement",  "Un departement (cours, EDT, notes, etudiants)",          "Admin Faculte"],
        ["professeur",         "Ses cours, TPs, notes, presences, messages",             "Admin Universite"],
        ["etudiant",           "Son espace personnel (lecture + soumissions)",           "Admin Dept / Auto-inscription"],
    ],
    col_widths=[4, 7.5, 4]
)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 3. ESPACE PUBLIC
# ════════════════════════════════════════════════════════════════════════════
h1(doc, "3. Espace Public (sans connexion)")

h2(doc, "3.1 Page d'accueil - 1_Accueil.py")
bullet(doc, "Affichage de toutes les universites actives avec logo, adresse et couleur personnalisee")
bullet(doc, "Statistiques temps reel : nb universites, professeurs, creneaux planifies, etudiants inscrits")
bullet(doc, "Recherche d'universite par nom ou adresse")
bullet(doc, "Cartes universites cliquables -> redirection vers l'emploi du temps")
bullet(doc, "Section CTA pour les institutions souhaitant rejoindre la plateforme")

h2(doc, "3.2 Emploi du temps public - 2_Horaire.py")
bullet(doc, "Selection en cascade : Universite -> Departement -> Promotion -> Classe")
bullet(doc, "Affichage grille hebdomadaire (Lundi au Samedi)")
bullet(doc, "Filtre semaines paires / impaires / toutes")
bullet(doc, "Filtre par type de creneau (Cours, TD, TP, Examen, Ferie)")
bullet(doc, "Affichage uniquement des creneaux dans leur periode de validite (valid_from / valid_until)")
bullet(doc, "Export iCal (.ics) pour synchronisation avec agenda personnel (Google, Outlook...)")
bullet(doc, "Affichage du professeur, de la salle, du code cours, du type de semaine")
bullet(doc, "Accessible sans connexion - lecture seule")
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 4. ESPACE ETUDIANT
# ════════════════════════════════════════════════════════════════════════════
h1(doc, "4. Espace Etudiant")

h2(doc, "4.1 Authentification - 10_Student_Auth.py")
bullet(doc, "Connexion par numero etudiant ou username + mot de passe")
bullet(doc, "Auto-inscription : l'etudiant cree son compte avec son numero (doit exister dans le registre)")
bullet(doc, "Verification du statut du compte : message specifique si compte desactive")

h2(doc, "4.2 Dashboard Etudiant - 11_Student_Dashboard.py")
add_table(doc,
    ["Onglet", "Fonctionnalites"],
    [
        ["Emploi du Temps",     "Grille hebdomadaire de sa classe. Filtre semaines paires/impaires. Export iCal. Creneaux expires masques automatiquement."],
        ["Cours & Documents",   "Telechargement des PDFs de cours deposes par les professeurs, filtres par matiere."],
        ["Mes TPs",             "Liste des TPs assignes. Depot de fichier PDF avant la deadline. Statut de soumission (depose / non depose). Fermeture automatique a la deadline."],
        ["Mes Notes",           "Notes publiees par matiere et type (examen / DS / TP / CC). Moyenne calculee. Soumission de reclamation si contestation."],
        ["Bulletin",            "Bulletin de notes complet par semestre. Export PDF. Statuts academiques (admis / redoublant / etc.)."],
        ["Annonces",            "Communiques de l'universite et du departement. Priorite aux annonces epinglees."],
        ["Messages",            "Messages recus des professeurs destines a sa classe."],
        ["Presences",           "Historique des presences / absences / retards enregistres par les professeurs."],
        ["Reclamations",        "Suivi des reclamations de notes soumises. Statut (en attente / accepte / refuse) et reponse du professeur."],
        ["Mon Compte",          "Changement de mot de passe (saisie ancien + nouveau x2). Affichage username et numero etudiant."],
    ],
    col_widths=[4, 11.5]
)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 5. ESPACE PROFESSEUR
# ════════════════════════════════════════════════════════════════════════════
h1(doc, "5. Espace Professeur")

h2(doc, "5.1 Authentification - 12_Prof_Auth.py")
bullet(doc, "Connexion par email + mot de passe")
bullet(doc, "Role verifie : acces refuse si le compte n'est pas de type 'professeur'")
bullet(doc, "Redirection automatique vers le dashboard si deja connecte")

h2(doc, "5.2 Dashboard Professeur - 9_Prof_Dashboard.py")
doc.add_paragraph(
    "Toutes les selections (Faculte, Departement, Promotion, Cours) sont filtrees "
    "dynamiquement : le professeur ne voit que les facultes/departements/cours "
    "auxquels il est reellement affecte (via horaires ou affectation directe sur le cours)."
)
add_table(doc,
    ["Onglet", "Fonctionnalites"],
    [
        ["Mon Emploi du Temps",
         "Grille personnelle par jour (Lundi-Samedi). Filtres : faculte, type de semaine, type de creneau. "
         "Chaque creneau affiche : horaire, type colore, cours, classe, promotion, departement, faculte, salle. "
         "Masquage automatique des creneaux dont la periode de validite est depassee."],
        ["Cours & Documents",
         "Depot de PDFs de cours avec cascade Faculte->Departement->Promotion->Cours. "
         "Liste de ses documents publies avec telechargement et suppression."],
        ["Travaux Pratiques",
         "Creation de TPs : titre, description, PDF sujet, dates debut et fin. "
         "Cascade de selection. Liste des TPs avec compteur de soumissions, "
         "telechargement des rendus etudiants, notation."],
        ["Notes",
         "Saisie et modification de notes par etudiant, cours, type (examen/DS/TP/CC), "
         "session (normale / rattrapage). Publication des notes. "
         "Gestion des demandes de modification de notes precedemment saisies."],
        ["Mes Classes",
         "Vue de toutes ses classes avec liste d'etudiants. "
         "Filtre par promotion / faculte via cascade."],
        ["Presences",
         "Marquage presence / absence / retard par seance et par etudiant. "
         "Cascade Faculte->Departement->Promotion->Classe."],
        ["Messages",
         "Envoi de messages a une classe entiere (avec option urgence). "
         "Cascade Faculte->Departement->Promotion->Classe. Hors du formulaire pour reactivite."],
        ["Reclamations",
         "Traitement des reclamations de notes : affichage note contestee, "
         "justification de l'etudiant, decision (accepter / refuser), reponse au texte libre. "
         "Filtre par faculte / departement / cours."],
    ],
    col_widths=[4, 11.5]
)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 6. ESPACE ADMINISTRATEUR
# ════════════════════════════════════════════════════════════════════════════
h1(doc, "6. Espace Administrateur - 8_Admin_Dashboard.py")
doc.add_paragraph(
    "Un seul fichier gere les 4 niveaux d'administration. "
    "L'interface s'adapte dynamiquement au role de l'admin connecte. "
    "Toutes les listes sont paginées (10 elements / page, 15 pour les etudiants)."
)

h2(doc, "6.1 Super Admin")
add_table(doc,
    ["Onglet", "Fonctionnalites"],
    [
        ["Universites",
         "Creer, modifier (nom, adresse, logo, couleur principale, site web). "
         "Desactiver / reactiver / supprimer definitivement. "
         "Stats par universite (nb facultes, departements, promotions). Pagination."],
        ["Comptes Administrateurs",
         "Creer des admins universite. Voir tous les comptes (actifs / inactifs). "
         "Modifier role, email, universite. Reinitialiser mot de passe. "
         "Desactiver / reactiver / supprimer. Pagination."],
        ["Analytiques",
         "Tableau de bord global : nb etudiants, profs, cours, creneaux par universite. "
         "Graphiques de repartition."],
    ],
    col_widths=[4, 11.5]
)

h2(doc, "6.2 Admin Universite")
add_table(doc,
    ["Onglet", "Fonctionnalites"],
    [
        ["Facultes & Departements",
         "Creer / modifier / supprimer des facultes. "
         "Ajouter des departements par faculte. Vue hierarchique paginee."],
        ["Professeurs",
         "Creer des profils professeurs (pool universite) : nom, email, telephone, statut contrat. "
         "Gestion complete du compte de connexion : creer (email + mdp), "
         "reinitialiser mdp, activer / desactiver le compte. "
         "Activer / desactiver le profil prof. Recherche par nom. Pagination."],
        ["Communiques",
         "Publier des annonces pour toute l'universite. "
         "Epingler, fixer une date d'expiration, joindre un fichier PDF ou image. Pagination."],
        ["Etudiants",
         "Vue lecture seule du registre etudiant complet. "
         "Filtres : annee academique, departement, statut, recherche nom/matricule. Pagination."],
        ["Admins",
         "Creer des admins faculte. Voir la liste des comptes admins de l'universite. Pagination."],
    ],
    col_widths=[4, 11.5]
)

h2(doc, "6.3 Admin Faculte")
add_table(doc,
    ["Onglet", "Fonctionnalites"],
    [
        ["Departements",
         "Creer, modifier, supprimer les departements de sa faculte. Pagination."],
        ["Professeurs",
         "Affilier un professeur du pool universite a sa faculte "
         "(statut : permanent = faculte principale, visiteur = enseigne en plus ici). "
         "Mettre a jour le statut d'affiliation. Retirer l'affiliation. Pagination."],
        ["Admins Departement",
         "Creer des comptes admins departement pour les departements de sa faculte."],
        ["Etudiants",
         "Vue lecture seule du registre etudiant de sa faculte. "
         "Filtres : annee, departement, statut, recherche. Pagination."],
    ],
    col_widths=[4, 11.5]
)

h2(doc, "6.4 Admin Departement")
add_table(doc,
    ["Onglet", "Fonctionnalites"],
    [
        ["Promotions & Groupes",
         "Creer / modifier des promotions (nom, annee academique, flag recrutement). "
         "Gerer les groupes/classes de chaque promotion (nom, capacite). Pagination."],
        ["Salles & Amphi",
         "Creer / modifier les salles : nom, code, capacite, type "
         "(amphi / labo / TP / salle info / salle / autre), batiment, etage. "
         "Filtre par type. Pagination."],
        ["Cours (EC)",
         "Creer / modifier les elements constitutifs : intitule, code, heures, "
         "coefficient, UE parente, promotion, professeur titulaire. Pagination."],
        ["Unites d'Enseignement (UE)",
         "Creer / modifier les UE : code, groupe (A-F), credits. "
         "Rattacher les EC aux UE. Pagination."],
        ["Professeurs",
         "Vue des professeurs affilies a la faculte du departement. "
         "Modification du profil (nom, email, telephone, statut contrat). Pagination."],
        ["Comptes Professeurs",
         "Vue lecture seule du statut des comptes de connexion des profs "
         "(compte actif / desactive / aucun compte). "
         "La creation et gestion des comptes se font chez l'admin universite."],
        ["Emploi du Temps",
         "Selection Promotion -> Classe. "
         "Creation de creneaux : jour, heure debut/fin, cours, professeur, salle, "
         "type (cours/TD/TP/examen/ferie), semaine (toutes/paire/impaire), dates validite, label libre. "
         "Modification et suppression. Export Excel et PDF. "
         "Historique complet des modifications de creneaux."],
        ["Communiques",
         "Publier des annonces pour le departement. "
         "Epingler, expiration, piece jointe. Pagination."],
        ["Notes & Resultats",
         "Tableau de bord des notes par promotion. "
         "Validation des demandes de modification de notes soumises par les profs. "
         "Calcul des moyennes et statuts academiques. Export."],
        ["Bulletin",
         "Generation et export PDF des bulletins de notes par etudiant, "
         "par semestre et annee academique."],
        ["Registre Etudiants",
         "Gestion complete des inscrits : ajouter manuellement, modifier statut academique "
         "(inscrit / admis / redoublant / transfere / abandonne), "
         "creer compte etudiant, import depuis fichier Excel. "
         "Filtres avances (annee, filiere, option, promotion, statut, recherche). "
         "Pagination cote serveur (50 par page)."],
        ["Reclamations",
         "Traiter les reclamations de notes en attente : "
         "voir note contestee, justification etudiant, accepter ou refuser avec reponse."],
        ["TPs & Soumissions",
         "Voir les soumissions de TPs par les etudiants. Telecharger les rendus. "
         "Filtre par cours / classe."],
    ],
    col_widths=[4.5, 11]
)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 7. BASE DE DONNEES
# ════════════════════════════════════════════════════════════════════════════
h1(doc, "7. Structure de la Base de Donnees")
doc.add_paragraph(
    "La base de donnees PostgreSQL compte 30 migrations (v1 a v30). "
    "Voici les tables principales avec leur role et leurs colonnes cles."
)
add_table(doc,
    ["Table", "Role", "Colonnes cles"],
    [
        ["universities",                   "Universites de la plateforme",                "id, name, primary_color, photo_url, is_active"],
        ["faculties",                       "Facultes d'une universite",                   "id, university_id, name, description"],
        ["departments",                     "Departements d'une faculte",                  "id, faculty_id, name, description"],
        ["promotions",                      "Annees d'etude d'un departement",             "id, department_id, name, academic_year, is_recrutement"],
        ["classes",                         "Groupes / sections d'une promotion",          "id, promotion_id, name, capacity"],
        ["professors",                      "Pool de professeurs d'une universite",        "id, university_id, name, email, phone, statut, is_active"],
        ["professor_faculty_affiliations",  "Affiliation prof -> faculte",                 "professor_id, faculty_id, status (permanent/visiteur)"],
        ["unites_enseignement",             "Unites d'enseignement (UE)",                  "id, department_id, code, name, credits, group_label"],
        ["courses",                         "Elements constitutifs (EC / matieres)",       "id, department_id, promotion_id, professor_id, ue_id, weight, hours"],
        ["schedules",                       "Creneaux emploi du temps",                    "id, class_id, course_id, professor_id, day, start_time, end_time, room_id, week_type, valid_from, valid_until, slot_type, slot_label"],
        ["rooms",                           "Salles et amphitheatres",                     "id, university_id, department_id, name, code, capacity, room_type, building"],
        ["users",                           "Comptes connexion (admins + profs)",          "id, email, role, password_hash, university_id, faculty_id, department_id, professor_id, is_active"],
        ["student_registry",                "Registre officiel des etudiants inscrits",   "id, student_number, university_id, class_id, statut, annee_academique"],
        ["students",                        "Comptes etudiants de connexion",              "id, student_number, registry_id, class_id, password_hash, username"],
        ["grades",                          "Notes des etudiants",                         "id, student_id, course_id, value, grade_type, session, is_published, professor_id"],
        ["grade_claims",                    "Reclamations de notes",                       "id, student_id, course_id, professor_id, status, response, faculty_id"],
        ["grade_modification_requests",     "Demandes de modif de notes (par prof)",       "id, professor_id, grade_id, new_value, reason, status"],
        ["grade_audit_log",                 "Historique modifications de notes",           "id, grade_id, old_value, new_value, changed_by, changed_at"],
        ["tp_assignments",                  "TPs crees par les professeurs",               "id, course_id, professor_id, class_id, title, description, deadline, is_active"],
        ["tp_submissions",                  "Rendus de TPs par les etudiants",             "id, assignment_id, student_id, file_url, submitted_at, grade"],
        ["course_documents",                "Documents de cours deposes (PDF)",            "id, course_id, professor_id, title, file_url, file_name, promotion_id"],
        ["attendance",                      "Presences / absences / retards",              "id, student_id, schedule_id, status, recorded_by, session_date"],
        ["class_messages",                  "Messages prof -> classe",                     "id, professor_id, class_id, content, is_urgent, created_at"],
        ["announcements",                   "Communiques universite / departement",        "id, university_id, department_id, title, content, is_pinned, expires_at, file_url"],
        ["bulletins",                       "Bulletins de notes generes",                  "id, student_id, semester, academic_year, pdf_url, generated_at"],
    ],
    col_widths=[4.5, 4.5, 6.5]
)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 8. FLUX UTILISATEUR
# ════════════════════════════════════════════════════════════════════════════
h1(doc, "8. Flux d'utilisation complet")

h2(doc, "8.1 Mise en service d'une universite (une seule fois)")
steps = [
    "Super Admin cree l'universite (nom, logo, couleur principale)",
    "Super Admin cree un compte Admin Universite",
    "Admin Universite cree les facultes",
    "Admin Universite cree les professeurs (profils + comptes de connexion email/mdp)",
    "Admin Faculte cree les departements",
    "Admin Faculte affilie les professeurs (permanent = principal, visiteur = secondaire)",
    "Admin Departement cree les UE et les EC (cours), les promotions, les classes et les salles",
    "Admin Departement assigne les professeurs titulaires aux cours",
    "Admin Departement construit les emplois du temps (creneaux par classe)",
    "Admin Departement importe ou cree les etudiants dans le registre",
    "Admin Departement cree les comptes etudiants (ou les etudiants s'auto-inscrivent)",
]
for i, s in enumerate(steps, 1):
    bullet(doc, "Etape " + str(i) + " : " + s)
doc.add_paragraph()

h2(doc, "8.2 Cycle academique hebdomadaire type")
steps2 = [
    "L'etudiant consulte son emploi du temps (public ou via son espace personnel)",
    "Le professeur depose les documents de cours (PDF) pour chaque matiere",
    "Le professeur cree les TPs avec description et deadline",
    "L'etudiant soumet son TP en PDF avant la deadline (fermeture automatique)",
    "Le professeur marque les presences en fin de seance",
    "Le professeur saisit et publie les notes",
    "L'etudiant consulte ses notes et peut soumettre une reclamation si contestation",
    "Le professeur traite la reclamation (accepte / refuse) avec reponse",
    "L'admin Departement valide les demandes de modification de notes",
    "L'admin Departement genere les bulletins en fin de semestre (export PDF)",
]
for i, s in enumerate(steps2, 1):
    bullet(doc, str(i) + ". " + s)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 9. SECURITE ET PERFORMANCES
# ════════════════════════════════════════════════════════════════════════════
h1(doc, "9. Securite & Performances")

h2(doc, "9.1 Securite implementee")
bullet(doc, "Mots de passe haches en bcrypt (cout adaptatif, irreversible)")
bullet(doc, "Sessions gerees via st.session_state (cote serveur Streamlit)")
bullet(doc, "Verification du role a chaque page (require_auth, require_prof_auth)")
bullet(doc, "Secrets DB et Supabase dans .streamlit/secrets.toml (exclu du depot Git)")
bullet(doc, "Requetes SQL preparees via psycopg2 (protection contre l'injection SQL)")
bullet(doc, "HTTPS obligatoire en production (certificat Let's Encrypt via Nginx)")
bullet(doc, "Isolation des donnees par universite (chaque admin ne voit que son perimetre)")

h2(doc, "9.2 Capacite de connexions simultanees")
add_table(doc,
    ["Configuration", "Connexions simultanees"],
    [
        ["Developpement local (maxconn=10)",            "10 - 20 utilisateurs"],
        ["VPS 2GB RAM + maxconn=25",                    "30 - 50 utilisateurs"],
        ["VPS 4GB RAM + maxconn=50 + cache Streamlit",  "50 - 150 utilisateurs"],
    ],
    col_widths=[9, 6.5]
)

h2(doc, "9.3 Architecture de deploiement recommandee")
bullet(doc, "VPS Hetzner CX21 (2 vCPU, 4GB RAM) - environ 5 euros/mois")
bullet(doc, "PostgreSQL heberge sur Supabase (gratuit jusqu'a 500MB)")
bullet(doc, "Streamlit gere par systemd (redemarrage automatique en cas de crash)")
bullet(doc, "Nginx en reverse proxy avec SSL Let's Encrypt (HTTPS gratuit)")
bullet(doc, "Supabase Storage pour tous les fichiers (PDFs, logos, pieces jointes)")
bullet(doc, "Sauvegardes automatiques via pg_dump quotidien + cron")
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 10. EVOLUTIONS POSSIBLES
# ════════════════════════════════════════════════════════════════════════════
h1(doc, "10. Evolutions possibles (Roadmap)")
add_table(doc,
    ["Fonctionnalite", "Priorite", "Complexite"],
    [
        ["Notifications email (note publiee, reclamation traitee, TP note)", "Haute",   "Moyenne"],
        ["Application mobile (PWA ou React Native)",                         "Haute",   "Haute"],
        ["Messagerie directe etudiant -> professeur",                        "Moyenne", "Faible"],
        ["Authentification OAuth (Google / Microsoft)",                      "Moyenne", "Moyenne"],
        ["Statistiques avancees (taux reussite par cours, assiduité)",       "Moyenne", "Moyenne"],
        ["Migration base de donnees vers Supabase PostgreSQL",               "Haute",   "Faible"],
        ["API REST publique pour integrations tierces (ERP universitaire)",  "Basse",   "Haute"],
        ["Multi-langue (francais / arabe / anglais)",                        "Basse",   "Moyenne"],
        ["Systeme de sondage / vote pour les etudiants",                     "Basse",   "Faible"],
        ["Tableau de bord analytique avance (graphiques interactifs)",       "Moyenne", "Moyenne"],
    ],
    col_widths=[8.5, 2.5, 2.5]
)

# ── Pied de page ──────────────────────────────────────────────────────────
doc.add_paragraph()
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_f = p_footer.add_run(
    "UniSchedule - Documentation generee le " + date_str + " - Version 2.0"
)
r_f.font.size = Pt(9)
r_f.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

doc.save("UniSchedule_Documentation.docx")
print("Document genere avec succes.")
